import speech_recognition as sr
import sounddevice as sd
import numpy as np
import wave
import threading
import time
import json
from datetime import datetime
from docx import Document
import google.generativeai as genai
from google.cloud import speech
import os
from scipy.io.wavfile import write

class MeetingRecorder:
    def __init__(self, api_key):
        self.is_recording = False
        self.audio_chunks = []
        self.sample_rate = 44100
        self.channels = 2
        
        # Configure Gemini API
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-pro')
        
        # Configure Speech-to-Text client
        self.speech_client = speech.SpeechClient()
        
    def start_recording(self):
        """Start recording the meeting audio"""
        self.is_recording = True
        self.recording_thread = threading.Thread(target=self._record_audio)
        self.recording_thread.start()
        print("Recording started...")

    def stop_recording(self):
        """Stop recording and process the audio"""
        self.is_recording = False
        self.recording_thread.join()
        print("Recording stopped.")
        
        # Save the recorded audio
        self._save_audio()
        
        # Transcribe the audio
        transcript = self._transcribe_audio()
        
        if transcript:
            # Generate summary
            summary = self._generate_summary(transcript)
            
            # Save transcript and summary to both .docx and .json
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            self._save_document(transcript, summary, timestamp)
            self._save_json(transcript, summary, timestamp)
        else:
            print("No transcript was generated. Please check your audio recording.")

    def _record_audio(self):
        """Record audio from the microphone"""
        with sd.InputStream(samplerate=self.sample_rate, channels=self.channels, callback=self._audio_callback):
            while self.is_recording:
                sd.sleep(100)

    def _audio_callback(self, indata, frames, time, status):
        """Callback function to store audio chunks"""
        if status:
            print(status)
        self.audio_chunks.append(indata.copy())

    def _save_audio(self):
        """Save recorded audio to a WAV file"""
        if not self.audio_chunks:
            return

        audio_data = np.concatenate(self.audio_chunks, axis=0)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"meeting_recording_{timestamp}.wav"
        
        # Save as 16-bit PCM WAV file
        write(filename, self.sample_rate, audio_data.astype(np.int16))
        self.latest_audio_file = filename
        print(f"Audio saved as {filename}")

    def _transcribe_audio(self):
        """Transcribe the audio using Google Cloud Speech-to-Text"""
        print("Transcribing audio...")
        
        try:
            # Read the audio file
            with open(self.latest_audio_file, 'rb') as audio_file:
                content = audio_file.read()
            
            # Configure the recognition
            audio = speech.RecognitionAudio(content=content)
            config = speech.RecognitionConfig(
                encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
                sample_rate_hertz=self.sample_rate,
                language_code="en-US",
                enable_automatic_punctuation=True,
            )
            
            # Perform the transcription
            response = self.speech_client.recognize(config=config, audio=audio)
            
            # Combine all transcriptions
            transcript = ""
            for result in response.results:
                transcript += result.alternatives[0].transcript + "\n"
            
            return transcript.strip()
            
        except Exception as e:
            print(f"Error during transcription: {e}")
            return ""

    def _generate_summary(self, transcript):
        """Generate a summary of the transcript using Google Gemini"""
        if not transcript:
            return "No transcript available to summarize."

        try:
            prompt = f"""Please provide a very concise summary (2-3 sentences) of the following meeting transcript:

{transcript}

Focus only on the most important points and any critical action items."""
            response = self.model.generate_content(prompt)
            return response.text
        except Exception as e:
            print(f"Error generating summary: {e}")
            return "Error generating summary."

    def _save_document(self, transcript, summary, timestamp):
        """Save the transcript and summary to a Word document"""
        doc = Document()
        
        # Add title
        doc.add_heading(f'Meeting Notes - {timestamp}', 0)
        
        # Add summary section
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)
        
        # Add transcript section
        doc.add_heading('Full Transcript', level=1)
        doc.add_paragraph(transcript)
        
        # Save document
        filename = f"meeting_notes_{timestamp}.docx"
        doc.save(filename)
        print(f"Document saved as {filename}")

    def _save_json(self, transcript, summary, timestamp):
        """Save the transcript and summary to a JSON file"""
        data = {
            'timestamp': timestamp,
            'summary': summary,
            'transcript': transcript
        }
        
        filename = f"meeting_notes_{timestamp}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        print(f"JSON file saved as {filename}")
